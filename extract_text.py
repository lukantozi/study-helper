import magic
import fitz 
import docx


def extract_from_txt(file_path):
    with open(file_path) as doc:
        text = doc.read()
    return text


def extract_from_pdf(file_path):
    doc = fitz.open(file_path)
    return '\n'.join([page.get_text() for page in doc])


def extract_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for p in doc.paragraphs:
        text += '\n' + p.text

    for table in doc.tables:
        text += '\n'
        for row in table.rows:
            text += '\n'
            for cell in row.cells:
                text += cell.text
                text += ","
    return text


def input_content():
    while True:
        mode = input("[T]ext or [F]ile: ").lower()
        if (mode == "text" or mode == "t"):
            mode = 1
            break
        elif (mode == "file" or mode == "f"):
            mode = 2
            break
    return mode


def extract_content():
    mode = input_content()
    if mode == 1:
        content = input("Please paste the content here: ")
        return content
    elif mode == 2:
        content = input("Please input the filepath here: ")

        try:
            content_type = magic.from_file(content, mime=True)
        except FileNotFoundError:
            return extract_content()

        if content_type == "text/plain":
            return extract_from_txt(content)
        elif content_type == "application/pdf":
            return extract_from_pdf(content)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_from_docx(content)
        else:
            print("Accepted formats: pdf, docx, txt")
            return extract_content()


def chunk():
    text = extract_content()
    n = 1000 # per question; switch to 10k-15k for 10-15 questions
    chunks = [text[i:i+n] for i in range(0, len(text), n)] # type: ignore
    return chunks


#text = extract_from_pdf("test.pdf")
#text = extract_from_docx("test.docx")
#text = extract_from_txt("test.txt")
#print(magic.from_file("test.docx", mime=True))
#text = extract_content()
#print(text)
#for chunk in chunks:
#    print(chunk)
