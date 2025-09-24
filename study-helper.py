import magic
import multiprocessing as mp
import subprocess
import string
from ollama import Client
from rake_nltk import Rake
import docx
import re
import json, csv, os, time
import unicodedata
try:
    import spacy
    _NLP = spacy.load("en_core_web_sm")
except Exception:
    _NLP = None
import math
from statistics import mean


# change the model when switching pc/laptop
OLLAMA_HOST = "http://127.0.0.1:11434"
LLM_MODEL   = "qwen2.5:7b-instruct"
EXTRA_DASHES = r"[–—−-]"  # en dash, em dash, minus, non-breaking hyphen
FORCE_RAKE = True


ollama_client = Client(host=OLLAMA_HOST)
# Fast preflight so we don't hang on generate() if the server isn't reachable
try:
    _ = ollama_client.list()  # pings /api/tags under the hood
except Exception as e:
    print(f"[ERR] Cannot reach Ollama at {OLLAMA_HOST}. Is the service running and listening on 127.0.0.1:11434?\n{e}")
    import sys
    sys.exit(1)

translator = str.maketrans('', '', string.punctuation)


def _safe_generate(**kwargs):
    try:
        return ollama_client.generate(**kwargs)
    except Exception as e:
        # surface and skip this item instead of hanging the whole run
        raise

def _strip_code_fences(t: str) -> str:
    # remove ```...``` wrappers if the model adds them
    t = t.strip()
    if t.startswith("```"):
        t = re.sub(r"^```.*?\n", "", t, flags=re.S)  # drop opening fence line
        t = re.sub(r"\n```$", "", t, flags=re.S)     # drop closing fence
    return t

def _canonicalize_tags(t: str) -> str:
    """
    Normalize all plausible tag variants to canonical 'Q:', 'A:', 'E:' at line start.
    Handles: Q1:, Question:, A1:, Answer:, Evidence:, e:, etc.  Case-insensitive.
    """
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # standardize common labels to canonical
    reps = [
        (r'(?mi)^\s*question\s*[:\-–]\s*', 'Q: '),
        (r'(?mi)^\s*q\d*\s*[:\-–]\s*',      'Q: '),
        (r'(?mi)^\s*answer\s*[:\-–]\s*',    'A: '),
        (r'(?mi)^\s*a\d*\s*[:\-–]\s*',      'A: '),
        (r'(?mi)^\s*evidence\s*[:\-–]\s*',  'E: '),
        (r'(?mi)^\s*e\d*\s*[:\-–]\s*',      'E: '),
    ]
    for pat, repl in reps:
        t = re.sub(pat, repl, t)
    return t.strip()

def parse_qae(text: str) -> dict:
    """
    Parse ONE Q/A/E triple from model output.
    Returns: {'ok': bool, 'q': str|None, 'a': str|None, 'e': str|None, 'raw': str, 'missing': set()}
    - 'a' can be multi-line; we stop at the next 'E:' line
    - 'e' is unquoted (quotes removed if present)
    """
    raw = _canonicalize_tags(_strip_code_fences(text))

    # Grab Q
    m_q = re.search(r'(?m)^Q:\s*(.*)$', raw)
    q = m_q.group(1).strip() if m_q else None

    # Grab A (multi-line, non-greedy, until next E:)
    m_a = re.search(r'(?ms)^A:\s*(.*?)\n\s*E:\s*', raw)
    a = m_a.group(1).strip() if m_a else None

    # Grab E (to end of string)
    m_e = re.search(r'(?m)^E:\s*(.*)$', raw)
    e = m_e.group(1).strip() if m_e else None
    if e:
        # remove surrounding quotes if present
        e = e.strip().strip('"“”„').strip()

    missing = set()
    if not q: missing.add('Q')
    if not a: missing.add('A')
    if not e: missing.add('E')

    return {'ok': len(missing) == 0, 'q': q, 'a': a, 'e': e, 'raw': raw, 'missing': missing}

def parse_qae_items(text: str) -> list[dict]:
    """
    Parse MULTIPLE Q/A/E items if the model returns more than one.
    Splits on 'Q:' lines and applies parse_qae to each block.
    Returns a list of dicts like parse_qae().
    """
    raw = _canonicalize_tags(_strip_code_fences(text))
    # split while keeping 'Q:' with the following line by using a lookahead
    blocks = re.split(r'(?m)^(?=Q:\s*)', raw)
    items = []
    for b in blocks:
        if b.strip():
            items.append(parse_qae(b))
    return items

def normalize_for_rake(s: str) -> str:
    s = re.sub(EXTRA_DASHES, "-", s)
    s = s.replace("\u00a0", " ")  # nbsp -> space
    s = s.replace("-", " ")
    # drop section markers like (1) (2) (3)
    s = re.sub(r"\(\d+\)\s*", "", s)
    # drop header lines like "Reading Material 1 – ..."
    s = re.sub(r"^reading material\s*\d*\s*[-–—]?\s*.*$", "", s, flags=re.I | re.M)
    # lowercase + collapse spaces
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s


def normalize_for_match(s: str) -> str:
    s = re.sub(EXTRA_DASHES, " ", s)
    s = s.replace("\u00a0", " ")
    s = s.replace("-", " ")
    s = s.lower()
    s = re.sub(r"[^\w\s]", " ", s)  # strip punctuation (ASCII + Unicode)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def extract_from_txt(file_path):
    with open(file_path) as doc:
        text = doc.read()
    return text


def extract_from_pdf_core(
    file_path,
    remove_headers=True,
    remove_footers=True,
    include_footnotes=True,
    detect_images=False,
    ocr_images=False,
    image_dir="extracted_images",
):
    """
    Robust text extraction with header/footer removal, reading-order sort,
    hyphen fix, and optional image OCR.

    Now hardened: any page-level extraction failure falls back to simpler
    text extraction so a single bad page can't crash the whole run.
    """
    import os
    import fitz  # PyMuPDF
    from statistics import median

    TOP_FRACTION = 0.08
    BOTTOM_FRACTION = 0.08
    FOOTNOTE_ZONE = 0.25
    REPEAT_THRESHOLD = 0.60
    Y_TOL = 4.0
    MIN_HEADER_LEN = 8
    MIN_FOOTER_LEN = 3

    # -------- open safely --------
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        # as a last resort, tell caller we couldn't parse
        raise RuntimeError(f"[PDF OPEN ERROR] {file_path}: {e}")

    try:
        pages_data = []
        top_candidates = []
        bot_candidates = []

        # First pass: gather candidate headers/footers
        for pno in range(len(doc)):
            page = doc[pno]
            width, height = page.rect.width, page.rect.height
            top_y = height * TOP_FRACTION
            bot_y = height * (1.0 - BOTTOM_FRACTION)

            # --- robust: if dict extraction crashes, skip to simple text ---
            try:
                d = page.get_text("dict")
            except Exception:
                # minimal fallback for header/footer detection
                d = {"blocks": []}

            lines = []
            for block in d.get("blocks", []):
                if block.get("type", 0) != 0:
                    continue
                for line in block.get("lines", []):
                    buff, sizes = [], []
                    for span in line.get("spans", []):
                        t = span.get("text", "")
                        if t:
                            buff.append(t)
                            try:
                                sizes.append(float(span.get("size", 0)))
                            except Exception:
                                pass
                    txt = "".join(buff).strip()
                    if not txt:
                        continue
                    bbox = line.get("bbox", [0, 0, 0, 0])
                    avg_size = (sum(sizes) / len(sizes)) if sizes else 0.0
                    lines.append((txt, bbox, avg_size))

            top_local, bot_local = [], []
            for txt, (x0, y0, x1, y1), _sz in lines:
                if y0 <= top_y and len(txt) >= MIN_HEADER_LEN:
                    norm = re.sub(r"\s+", " ", txt).strip()
                    top_local.append(norm)
                elif y1 >= bot_y and len(txt) >= MIN_FOOTER_LEN:
                    norm = re.sub(r"\s+", " ", txt).strip()
                    bot_local.append(norm)

            top_candidates.append(set(top_local))
            bot_candidates.append(set(bot_local))

            pages_data.append({
                "size": (width, height),
                "page": page,
            })

        def decide_repeating(candidates_list, threshold, page_count):
            from collections import Counter
            c = Counter()
            for s in candidates_list:
                c.update(s)
            min_pages = int(page_count * threshold + 0.5)
            return set([s for s, n in c.items() if n >= min_pages])

        headers = decide_repeating(top_candidates, REPEAT_THRESHOLD, len(doc)) if remove_headers else set()
        footers = decide_repeating(bot_candidates, REPEAT_THRESHOLD, len(doc)) if remove_footers else set()

        body_parts, footnotes_all = [], []

        if detect_images:
            os.makedirs(image_dir, exist_ok=True)

        # Second pass: extract body/footnotes with per-page fallback
        for pno, pdata in enumerate(pages_data):
            page = pdata["page"]
            width, height = pdata["size"]

            # Optional image export (wrapped)
            if detect_images:
                try:
                    for img in page.get_images(full=True):
                        xref = img[0]
                        try:
                            pix = page.get_pixmap(xref=xref)
                            out = os.path.join(image_dir, f"p{pno+1}_x{xref}.png")
                            pix.save(out)
                        except Exception:
                            try:
                                dimg = doc.extract_image(xref)
                                ext = dimg.get("ext", "png")
                                img_bytes = dimg.get("image")
                                with open(os.path.join(image_dir, f"p{pno+1}_x{xref}.{ext}"), "wb") as f:
                                    f.write(img_bytes)
                            except Exception:
                                pass
                except Exception:
                    pass

            # Try rich layout first
            lines = []
            try:
                d = page.get_text("dict")
                for block in d.get("blocks", []):
                    if block.get("type", 0) != 0:
                        continue
                    for line in block.get("lines", []):
                        buff, sizes = [], []
                        for span in line.get("spans", []):
                            t = span.get("text", "")
                            if t:
                                buff.append(t)
                                try:
                                    sizes.append(float(span.get("size", 0)))
                                except Exception:
                                    pass
                        txt = "".join(buff).strip()
                        if not txt:
                            continue
                        bbox = line.get("bbox", [0, 0, 0, 0])
                        avg_size = (sum(sizes) / len(sizes)) if sizes else 0.0
                        lines.append((txt, bbox, avg_size))
            except Exception:
                # Fallback: plain text only
                try:
                    plain = page.get_text("text") or ""
                    if plain.strip():
                        body_parts.append(plain.strip())
                    continue
                except Exception:
                    # Worst-case: skip this page
                    continue

            # sort lines (layout-aware)
            def sort_key(item):
                _txt, (x0, y0, x1, y1), _sz = item
                return (round(y0 / Y_TOL), x0)
            lines.sort(key=sort_key)

            page_body, page_foot = [], []
            for txt, (x0, y0, x1, y1), sz in lines:
                norm = re.sub(r"\s+", " ", txt).strip()
                if remove_headers and norm in headers:
                    continue
                if remove_footers and norm in footers:
                    continue
                # drop bare page numbers close to margins
                if (y0 <= height * TOP_FRACTION or y1 >= height * (1.0 - BOTTOM_FRACTION)) and re.fullmatch(r"\d{1,4}", norm):
                    continue

                if include_footnotes and (y1 >= height * (1.0 - FOOTNOTE_ZONE)):
                    page_foot.append((txt, sz))
                else:
                    page_body.append(txt)

            page_text = "\n".join(page_body)
            page_text = re.sub(r"(\w)-\n(\w)", r"\1\2", page_text)
            page_text = re.sub(r"[ \t]+\n", "\n", page_text)
            page_text = re.sub(r"\n{3,}", "\n\n", page_text)

            if page_text.strip():
                body_parts.append(page_text.strip())

            if include_footnotes and page_foot:
                sizes = [sz for _t, sz in page_foot if sz > 0]
                cutoff = (median(sizes) * 0.8) if sizes else 8.0
                notes = [t for t, sz in page_foot if sz > 0 and sz <= cutoff]
                if notes:
                    footnotes_all.extend([f"[p.{pno+1}] {n}" for n in notes])

        # Optional OCR (kept but wrapped)
        if detect_images and ocr_images:
            try:
                import pytesseract
                from PIL import Image
                ocr_texts = []
                for fn in sorted(os.listdir(image_dir)):
                    if fn.lower().endswith((".png", ".jpg", ".jpeg", ".tif", ".bmp", ".webp")):
                        ocr_texts.append(pytesseract.image_to_string(Image.open(os.path.join(image_dir, fn))))
                if ocr_texts:
                    body_parts.append("\n\n--- OCR (images) ---\n" + "\n\n".join(t.strip() for t in ocr_texts if t.strip()))
            except Exception:
                pass

        final_text = "\n\n".join(bp for bp in body_parts if bp)
        if include_footnotes and footnotes_all:
            final_text += "\n\n--- FOOTNOTES ---\n" + "\n".join(footnotes_all)
        return final_text

    finally:
        try:
            doc.close()
        except Exception:
            pass


def extract_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for p in doc.paragraphs:
        text += '\n' + p.text

    for table in doc.tables:
        text += '\n'
        for row in table.rows:
            text += '\n'
            for cell in row.cells:
                text += cell.text
                text += ","
    return text


CTRL_RE = re.compile(r'[\x00-\x1F\x7F]')

def clean_text(s: str) -> str:
    if not s: return ""
    s = unicodedata.normalize('NFKC', s)
    # common ligatures / punctuation / artifacts
    s = (s.replace('ﬂ','fl')
           .replace('ﬁ','fi')
           .replace('ﬀ','ff')
           .replace('ﬃ','ffi')
           .replace('ﬄ','ffl')
           .replace('—','-').replace('–','-').replace('−','-')
           .replace('“','"').replace('”','"').replace("’","'")
           .replace('×','x'))  # or '×' -> 'x' for consistency
    s = CTRL_RE.sub('', s)
    return s


def _fitz_child_worker(path, kwargs, q):
    """
    Runs MuPDF extraction in a separate process so crashes don't kill the parent.
    Put result in a queue; on failure, put ''.
    """
    try:
        text = extract_from_pdf_core(path, **kwargs)
        q.put(text or "")
    except Exception:
        q.put("")

def _fallback_pdfminer(path: str) -> str:
    try:
        from pdfminer_high_level import extract_text as pdfminer_extract_text  # may differ in envs
    except Exception:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract_text
        except Exception:
            return ""
    try:
        return pdfminer_extract_text(path) or ""
    except Exception:
        return ""

def _fallback_pypdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n\n".join((p.extract_text() or "") for p in reader.pages) or ""
    except Exception:
        return ""

def _fallback_pdftotext_cli(path: str) -> str:
    # optional: use Poppler CLI if installed
    try:
        out = subprocess.check_output(["pdftotext", "-layout", path, "-"], stderr=subprocess.DEVNULL)
        return out.decode("utf-8", errors="replace")
    except Exception:
        return ""

def extract_from_pdf_safe(path: str, timeout_sec: int = 45, **kwargs) -> str:
    """
    Safe front door for PDF extraction.
    1) Spawn a child process to run MuPDF.
    2) If child crashes/hangs, terminate and fall back to pdfminer/pypdf/CLI.
    """
    q = mp.Queue()
    p = mp.Process(target=_fitz_child_worker, args=(path, kwargs, q))
    p.start()
    p.join(timeout_sec)

    text = ""
    if p.is_alive():
        # hung -> terminate
        try: p.terminate()
        except Exception: pass
        p.join(2)

    if p.exitcode == 0:
        try:
            text = q.get_nowait()
        except Exception:
            text = ""

    if text.strip():
        print("[INFO] PDF text via: MuPDF (child)")
        return text

    # --- fallbacks ---
    text = _fallback_pdfminer(path)
    if text.strip():
        print("[INFO] PDF text via: pdfminer.six")
        return text

    text = _fallback_pypdf(path)
    if text.strip():
        print("[INFO] PDF text via: pypdf")
        return text

    text = _fallback_pdftotext_cli(path)
    print("[INFO] PDF text via: pdftotext CLI")
    return text


def input_content():
    while True:
        mode = input("[T]ext or [F]ile: ").lower()
        if (mode == "text" or mode == "t"):
            mode = 1
            break
        elif (mode == "file" or mode == "f"):
            mode = 2
            break
    return mode


def extract_content(mode):
    if mode == 1:
        return input("Please paste the content here: ")

    # mode == 2
    path = input("Please input the filepath here: ").strip()
    if not os.path.exists(path):
        print(f"[ERR] File not found: {path}")
        return extract_content(mode)

    # Show the absolute path so you know exactly what you're opening
    abs_path = os.path.abspath(path)
    print(f"[INFO] Opening: {abs_path}")

    try:
        ctype = magic.from_file(abs_path, mime=True)
        print(f"[INFO] Detected MIME type: {ctype}")
    except Exception as e:
        print(f"[WARN] libmagic failed on {abs_path}: {e}")
        # Fallback by extension
        ext = os.path.splitext(abs_path)[1].lower()
        if ext in {".txt"}:
            ctype = "text/plain"
        elif ext in {".pdf"}:
            ctype = "application/pdf"
        elif ext in {".docx"}:
            ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            print(f"[ERR] Unsupported type by extension: {ext} (accepted: pdf, docx, txt)")
            return extract_content(mode)

    if ctype == "text/plain":
        with open(abs_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    if ctype == "application/pdf":
        return extract_from_pdf_safe(abs_path)

    if ctype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_from_docx(abs_path)

    print(f"[ERR] Unsupported type: {ctype} (accepted: pdf, docx, txt)")
    return extract_content(mode)




def split_sentences(text):
    # split on sentence end *or* line breaks
    return [s.strip() for s in re.split(r'(?<=[.!?])\s+|\n+', text) if s.strip()]


def chunks(chunk_size, sentences):
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 <= chunk_size:
            current_chunk += " " + sentence.strip()
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence.strip()
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks


def join_chunks(
    chunk_size: int,
    text: str,
    overlap_sents: int = 0,
    min_chunk_chars: int = 1800,
    max_chunk_chars: int = 3200
):
    sents = split_sentences(text)
    if not sents:
        return []

    chunks, cur = [], []
    cur_len = 0
    i = 0
    while i < len(sents):
        s = sents[i].strip()
        if not s:
            i += 1
            continue

        # default target is chunk_size, but we also enforce min_chunk_chars
        next_len = cur_len + len(s) + 1
        if next_len <= chunk_size or (cur_len < min_chunk_chars and next_len <= max_chunk_chars):
            cur.append(s)
            cur_len = next_len
            i += 1
            continue

        # time to emit; if still too small, force-grow up to max_chunk_chars
        if cur and cur_len < min_chunk_chars and i < len(sents):
            while i < len(sents) and cur_len < min_chunk_chars and cur_len + len(sents[i]) + 1 <= max_chunk_chars:
                cur.append(sents[i].strip())
                cur_len += len(sents[i]) + 1
                i += 1

        if cur:
            chunks.append(" ".join(cur).strip())

            # overlap tail (if requested)
            tail = cur[-overlap_sents:] if overlap_sents > 0 else []
            cur = tail[:]
            cur_len = sum(len(x) + 1 for x in cur)
        else:
            # pathological single long sentence
            chunks.append(s[:max_chunk_chars])
            i += 1
            cur, cur_len = [], 0

    if cur:
        chunks.append(" ".join(cur).strip())

    return chunks


def _sent_stats(text: str):
    sents = split_sentences(text)
    if not sents:
        return 0, 0, 0
    lens = [len(s) for s in sents]
    avg = mean(lens)
    p50 = sorted(lens)[len(lens)//2]
    p90 = sorted(lens)[int(len(lens)*0.9)]
    return avg, p50, p90

def choose_chunking(text: str):
    """
    Heuristic chunking:
    - Aim for ~ 12–24 sentences per chunk depending on sentence length.
    - Never go below 1800 chars or above 3200 chars.
    - Drop overlap if sentences are long; add a bit if short/clipped.
    """
    N = len(text)
    avg, p50, p90 = _sent_stats(text)

    # Base size from sentence statistics
    if p50 == 0:
        target = 2400
    else:
        target = int(min(max(p50 * 40, 1800), 3200))

    # Adjust by document size bands
    if N < 8_000:
        target = max(target, 2200)
    elif N > 80_000:
        target = min(target, 2000)

    # Overlap
    if p90 > 180:
        overlap = 0
    elif p50 > 100:
        overlap = 0
    elif p50 > 60:
        overlap = 1
    else:
        overlap = 2

    target = int(min(max(target, 1800), 3200))
    return target, overlap


FEW_SHOT = ""

def build_qg_prompt(chunk_text: str, a) -> str:
    # <<< FIX: relaxed, natural prompt (no “what it is, one property” cage)
    return (
        "Write one high-quality study Q/A/E from THIS CHUNK ONLY.\n"
        "Language: English.\n"
        f'Anchor phrase: "{a}" (must appear verbatim in the question)\n\n'
        "Output exactly:\n"
        "Q: <one open question>\n"
        "A: <a clear answer in 2–4 sentences based only on the chunk>\n"
        'E: "<copy exactly one supporting sentence from the chunk, with the period>"\n\n'
        "Constraints:\n"
        "• Start Q with How / Why / Explain / Compare / Under what conditions.\n"
        "• Do NOT use phrases like 'this chunk/passage/text'.\n"
        "• No outside facts.\n\n"
        f"Chunk:\n{chunk_text}"
    )


def build_reg_prompt(chunk_text: str, a, q) -> str:
    return (
        "Use only this chunk.\n"
        "Respond in English only.\n"
        f'Regenerate this question about "{a}": {q}\n'
        "Output exactly:\n"
        "Q: <one question>\n"
        "A: <3–4 sentences that cover the core idea and one concrete detail from the chunk that is not your E sentence>\n"
        'E: "<one exact sentence from the chunk; include the final period>"\n'
        "Rules:\n"
        f"• Include the anchor phrase verbatim: {a}\n"
        "• Prefer How/Why/Explain/Compare. No “in the chunk/text”.\n"
        "• Do not invent facts.\n\n"
        f"Chunk:\n{chunk_text}"
    )

def _gen_worker(args, q):
    try:
        resp = ollama_client.generate(**args)
        q.put(resp.get("response", ""))
    except Exception:
        q.put("")

def safe_llm_call(args, timeout_sec=25):
    q = mp.Queue()
    p = mp.Process(target=_gen_worker, args=(args, q))
    p.start()
    p.join(timeout_sec)
    if p.is_alive():
        try: p.terminate()
        except Exception: pass
        p.join(2)
        return ""  # timed out -> treat as failure
    try:
        return q.get_nowait()
    except Exception:
        return ""


def _first_sentence_with_anchor(chunk_text: str, anchor: str) -> str:
    sents = re.split(r'(?<=[.!?])\s+', chunk_text.strip())
    n_anchor = normalize_for_match(anchor)
    for s in sents:
        if n_anchor in normalize_for_match(s):
            return s.strip()
    return sents[0].strip() if sents else chunk_text.strip()

def llm_generate_questions(chunk_text: str, anchor, q=None, temperature: float = 0.2) -> str:
    opts = {
        "temperature": 0.4,
        "top_p": 0.95,
        "num_predict": 256,
        "gpu_layers": 999,
        "num_ctx": 1024,
        "num_batch": 128,
    }
    prompt = build_reg_prompt(chunk_text, anchor, q) if q else build_qg_prompt(chunk_text, anchor)
    args = {"model": LLM_MODEL, "prompt": prompt, "options": opts}
    resp = safe_llm_call(args, timeout_sec=25)
    if not resp.strip():
        # <<< FIX: anchor-aware, natural How/Why fallback (no “Explain … what it is …”)
        e = _first_sentence_with_anchor(chunk_text, anchor)
        if not re.search(r'[.!?]"?$', e): e += '.'
        return (
            f"Q: How does {anchor} influence the ideas presented, and why does it matter?\n"
            "A: It plays a central role within the discussion by shaping how concepts connect and what results are possible. "
            "Based on the passage, its use affects the structure/behavior described and leads to practical consequences for the scenario. "
            "These consequences explain why the concept is emphasized and how it guides reasoning in this part of the text.\n"
            f'E: "{e}"'
        )
    return resp


def get_qae_or_regen(chunk_raw: str, anchor_raw: str, prev_q: str | None = None):
    # 1st attempt
    resp = llm_generate_questions(chunk_raw, anchor_raw, q=prev_q, temperature=0.2)
    item = parse_qae(resp)

    # If malformed (missing Q/A/E), try one regen
    if not item['ok']:
        resp = llm_generate_questions(chunk_raw, anchor_raw, q=(item.get('q') or prev_q or ""), temperature=0.2)
        item = parse_qae(resp)

    return item  # dict with keys: ok, q, a, e, missing


BAD_ANCHOR_PAT = re.compile(
    r"""(?ix)
    ^(?:figure|fig|table|chapter|section|page|pp?|http|www)\b
    |^\d{1,4}$
    |^\w{1,2}$
    |^[^A-Za-z]*$
    |(?:mcs\s*-\s*ftl|mcs\s*ftl)
    """
)
COMMON_STOP_PHRASES = set("""
abstract introduction conclusion references appendix acknowledgments copyright
figure fig table chapter section page pages url http https www et al
problem example exercise theorem lemma proof definition remark note
year years january february march april may june july august september october november december
""".split())

def _sentences(s: str) -> list[str]:
    return [x for x in re.split(r'(?<=[.!?])\s+', s) if x.strip()]

def _token_words(s: str) -> list[str]:
    return re.findall(r"[A-Za-z][A-Za-z0-9\-']+", s)

def _is_stop_phrase(s: str) -> bool:
    w = normalize_for_match(s)
    return w in COMMON_STOP_PHRASES or BAD_ANCHOR_PAT.search(w) is not None

def _score_anchor(s: str) -> float:
    score = 0.0
    s0 = s.strip()
    if 3 <= len(s0) <= 70 and re.search(r"[A-Za-z]", s0) and not BAD_ANCHOR_PAT.search(s0):
        score += 1.0
    if " " in s0:
        score += 0.3
    if re.search(r"[0-9/\\|]", s0):
        score -= 0.5
    return score

def extract_keywords(chunk: str) -> list[str]:
    c = re.sub(r"\s+", " ", chunk.strip())
    cand = []
    if _NLP and not FORCE_RAKE:
        doc = _NLP(c)
        for nc in doc.noun_chunks:
            t = re.sub(EXTRA_DASHES, "-", nc.text.strip())
            cand.append(t)
    else:
        r = Rake(min_length=2, max_length=5)
        r.extract_keywords_from_text(normalize_for_rake(c))
        cand = r.get_ranked_phrases()

    seen, cleaned = set(), []
    for s in cand:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            cleaned.append(s)

    # score & filter
    scored = [(s, _score_anchor(s)) for s in cleaned]
    scored = [s for (s, sc) in scored if sc > 0]

    # If RAKE/NLP failed, derive anchors from the chunk itself
    if not scored:
        tokens = re.findall(r"[A-Za-z]{3,}", c.lower())
        stop = set("""
            the a an and or of in on to for with by from as at is are was were be been being
            this that these those it its their his her our your my we you they them he she
            into over under between among about across within without during before after
            not no yes more most less least very just also
        """.split())
        freq = {}
        for t in tokens:
            if t in stop:
                continue
            freq[t] = freq.get(t, 0) + 1
        bigrams = []
        for i in range(len(tokens) - 1):
            a, b = tokens[i], tokens[i+1]
            if a in stop or b in stop:
                continue
            bigrams.append(f"{a} {b}")
        bigf = {}
        for bg in bigrams:
            bigf[bg] = bigf.get(bg, 0) + 1

        top_bi = sorted(bigf.items(), key=lambda x: x[1], reverse=True)[:8]
        top_uni = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:8]
        fallback = [k for k,_ in top_bi] + [k for k,_ in top_uni]

        def ok_anchor(s):
            return 3 <= len(s) <= 70 and re.search(r"[a-z]", s) and not BAD_ANCHOR_PAT.search(s)

        fallback = [s for s in fallback if ok_anchor(s)]
        if not fallback:
            fallback = ["key idea"]

        return fallback[:8]

    # normal path: rank by frequency & length
    low = normalize_for_match(chunk)
    freq = {}
    for s in list(scored):
        n = normalize_for_match(s)
        freq[s] = low.count(n)
    scored.sort(key=lambda s: (freq.get(s, 0), len(s)), reverse=True)
    return scored[:8]


BAD_SINGLE_WORDS = {
    "use","way","one","two","three","form","word","think","learn","plan",
    "different","exercise","example","problem","figure","table","section",
    "reading","material","device","devices","technology","technologies"
}
BAD_CHARS = set("•/\\|,;:=+[]{}()<>«»·•\u00b7\u2022")

def _alpha_ratio(s: str) -> float:
    a = sum(ch.isalpha() for ch in s)
    return a / max(1, len(s))

def clean_anchors(cands: list[str], chunk_raw: str, want_n: int = 8) -> list[str]:
    out = []
    seen = set()
    chunk_norm = normalize_for_match(chunk_raw)

    for s in cands:
        s0 = s.strip()
        if not s0:
            continue
        key = s0.lower()
        if key in seen:
            continue
        seen.add(key)

        # drop if weird punctuation/bullets present
        if any(ch in BAD_CHARS for ch in s0):
            continue
        if _alpha_ratio(s0) < 0.65:
            continue

        toks = re.findall(r"[A-Za-z][A-Za-z0-9\-']+", s0)
        if len(toks) == 0:
            continue
        # prefer multi-word anchors; allow single only if not generic
        if len(toks) == 1 and toks[0].lower() in BAD_SINGLE_WORDS:
            continue

        # length sanity
        phrase = " ".join(toks)
        if len(phrase) < 4 or len(phrase) > 70:
            continue

        # must actually occur (normalized) in the chunk
        if normalize_for_match(s0) not in chunk_norm:
            continue

        out.append(s0)

    # rank: multiword first, then by frequency in chunk, then by length
    def rank_key(s: str):
        toks = re.findall(r"[A-Za-z][A-Za-z0-9\-']+", s)
        freq = chunk_norm.count(normalize_for_match(s))
        return (
            0 if len(toks) >= 2 else 1,
            -freq,
            -len(s)
        )

    out.sort(key=rank_key)

    # fallback if we filtered everything: pick frequent bi/tri-grams
    if not out:
        words = re.findall(r"[A-Za-z][A-Za-z0-9']+", chunk_raw.lower())
        grams: dict[str,int] = {}
        for n in (3, 2):
            for i in range(len(words) - n + 1):
                g = " ".join(words[i:i+n])
                if _alpha_ratio(g) < 0.9:
                    continue
                if any(w in BAD_SINGLE_WORDS for w in g.split()):
                    continue
                grams[g] = grams.get(g, 0) + 1

        if grams:
            out = [g for g, _cnt in sorted(grams.items(), key=lambda kv: kv[1], reverse=True)][:max(3, want_n // 2)]
        else:
            out = ["key idea"]

    return out[:want_n]


def dedupe_anchors(anchors: list[str]) -> list[str]:
    # <<< FIX: simple, non-destructive dedupe (no aggressive lemmatization)
    out, seen = [], set()
    for a in anchors:
        key = normalize_for_match(a)
        if key in seen:
            continue
        seen.add(key)
        out.append(a.strip())
    return out



def pick_anchor(kw_8: list[str], chunk_raw: str) -> str:
    if not chunk_raw.strip():
        return kw_8[0] if kw_8 else "key idea"

    low_chunk = normalize_for_match(chunk_raw)
    sents_norm = [normalize_for_match(s) for s in _sentences(chunk_raw)]

    def score_candidate(cand: str) -> float:
        if not cand or _is_stop_phrase(cand):
            return -1e9
        n_cand = normalize_for_match(cand)
        if not re.search(r"[a-z]", n_cand):
            return -1e9
        freq = low_chunk.count(n_cand)
        if freq == 0:
            return -1e9
        coverage = sum(1 for sn in sents_norm if n_cand in sn)
        tok_count = len(_token_words(cand))
        letter_ratio = len(re.findall(r"[A-Za-z]", cand)) / max(1, len(cand))
        has_digit = bool(re.search(r"\d", cand))
        multiword_bonus = 0.4 if tok_count >= 2 else 0.0
        length_bonus = min(tok_count, 6) * 0.1
        digit_penalty = 0.6 if has_digit else 0.0
        header_penalty = 0.8 if BAD_ANCHOR_PAT.search(cand) else 0.0
        stop_like = 0.5 if cand.lower() in COMMON_STOP_PHRASES else 0.0

        return (2.0 * freq) + (1.2 * coverage) + multiword_bonus + length_bonus \
               + (0.3 * letter_ratio) - digit_penalty - header_penalty - stop_like

    scored = [(cand, score_candidate(cand)) for cand in kw_8]
    scored = [x for x in scored if x[1] > -1e8]
    if scored:
        scored.sort(key=lambda t: t[1], reverse=True)
        return scored[0][0]

    words = _token_words(chunk_raw)
    grams = set()
    for n in (3, 2):
        for i in range(len(words) - n + 1):
            cand = " ".join(words[i:i+n])
            if not _is_stop_phrase(cand):
                grams.add(cand)

    scored_fallback = [(g, score_candidate(g)) for g in grams]
    scored_fallback = [x for x in scored_fallback if x[1] > -1e8]
    if scored_fallback:
        scored_fallback.sort(key=lambda t: t[1], reverse=True)
        return scored_fallback[0][0]

    return kw_8[0] if kw_8 else "key idea"


def _sentence_count(s: str) -> int:
    return len([x for x in re.split(r'[.!?](?:\s+|$)', s.strip()) if x.strip()])

def quality_ok(anchor_raw, q, a, e, chunk_raw) -> bool:
    if not q or not a or not e:
        return False
    if not re.match(r'^(How|Why|Explain|Compare|Under what conditions)\b', q):
        return False
    if normalize_for_match(anchor_raw) not in normalize_for_match(q):
        return False
    if normalize_for_match(anchor_raw) not in normalize_for_match(e):
        return False
    if normalize_for_match(e) not in normalize_for_match(chunk_raw):
        return False
    sc = _sentence_count(a)
    if sc < 2 or sc > 5:
        return False
    if BAD_ANCHOR_PAT.search(q.lower()):
        return False
    return True

def extract_q_e(text):
    q = text.split("Q:", 1)[1].split("\n", 1)[0].strip()
    e = text.split('E:', 1)[1].strip()
    # strip surrounding quotes if present
    if e.startswith('"') and '"' in e[1:]:
        e = e[1:e[1:].find('"')+1]  # keep inside quotes incl. closing quote
        e = e.strip('"').strip()
    return e, q
        

def validate_qs(chunk, keywords, evidence):
    ch = normalize_for_match(chunk)
    ev = normalize_for_match(evidence)
    kws = [normalize_for_match(kw) for kw in keywords]
    return ch, kws, ev
    

def check_evidence_chunk(chunk_norm, chunk_raw, anchor_norm, anchor_raw, e_raw, q_raw, a_raw, kw_8, regen_left=1):
    """
    1) Evidence must be verbatim substring of RAW chunk.
    2) If not, try one regen (parse with parse_qae) and re-check.
    3) If still bad and regen_left exhausted, return what we have.
    """
    if e_raw and e_raw in chunk_raw:
        return check_anchor_evidence(chunk_norm, chunk_raw, anchor_norm, anchor_raw, e_raw, q_raw, a_raw, kw_8, regen_left)

    if regen_left <= 0:
        return q_raw, a_raw, e_raw

    resp = llm_generate_questions(chunk_raw, anchor_raw, q=q_raw, temperature=0.4)
    item = parse_qae(resp)
    if not item['ok']:
        return q_raw, a_raw, e_raw

    return check_evidence_chunk(
        chunk_norm, chunk_raw, anchor_norm, anchor_raw,
        item['e'], item['q'], item['a'], kw_8, regen_left=0
    )


def check_anchor_evidence(chunk_norm, chunk_raw, anchor_norm, anchor_raw, e_raw, q_raw, a_raw, kw_8, regen_left):
    """
    1) Anchor must appear (normalized) inside the evidence.
    2) If not, try one regen (parse) with same anchor.
    3) If still bad, try the next anchor from kw_8 (fresh Q/A/E).
    """
    if e_raw and anchor_norm in normalize_for_match(e_raw):
        return q_raw, a_raw, e_raw

    if regen_left > 0:
        resp = llm_generate_questions(chunk_raw, anchor_raw, q=q_raw, temperature=0.4)
        item = parse_qae(resp)
        if item['ok'] and anchor_norm in normalize_for_match(item['e']):
            return item['q'], item['a'], item['e']

    # Try next anchor
    try:
        i = kw_8.index(anchor_raw)
        next_anchor_raw = kw_8[i+1]
    except (ValueError, IndexError):
        return q_raw, a_raw, e_raw

    next_anchor_norm = normalize_for_match(next_anchor_raw)

    resp = llm_generate_questions(chunk_raw, next_anchor_raw, q=None, temperature=0.2)
    item = parse_qae(resp)
    if not item['ok']:
        return q_raw, a_raw, e_raw

    return check_evidence_chunk(
        chunk_norm, chunk_raw, next_anchor_norm, next_anchor_raw,
        item['e'], item['q'], item['a'], kw_8, regen_left=1
    )

def _norm_q(q: str) -> str:
    return normalize_for_match(q)

def _tokens(s: str) -> set[str]:
    return set(re.findall(r"[A-Za-z][A-Za-z0-9\-']+", s.lower()))

def _near_dupe(q: str, seen_norm_qs: set[str], seen_token_sets: list[set[str]], jaccard_threshold: float = 0.75) -> bool:
    nq = _norm_q(q)
    if nq in seen_norm_qs:
        return True
    qt = _tokens(q)
    if not qt:
        return False
    for st in seen_token_sets:
        inter = len(qt & st)
        if inter == 0:
            continue
        jac = inter / float(len(qt | st))
        if jac >= jaccard_threshold:
            return True
    return False


def main():
    mode = input_content()
    text = clean_text(extract_content(mode))

    print(f"[INFO] Extracted chars: {len(text)}")

    MAX_CHARS = 200_000
    if len(text) > MAX_CHARS:
        print(f"[WARN] Truncating text to {MAX_CHARS} chars for safety.")
        text = text[:MAX_CHARS]

    # --- automatic chunking ---
    size, overlap = choose_chunking(text)
    chunks = join_chunks(size, text, overlap_sents=overlap)
    print(f"[INFO] Num chunks: {len(chunks)} (target={size}, overlap={overlap})")

    # If we exploded into too many chunks, increase size until safe
    MAX_CHUNKS = 60
    while len(chunks) > MAX_CHUNKS and size < 3600:
        size = int(size * 1.2)
        size = min(size, 3600)
        chunks = join_chunks(size, text, overlap_sents=max(0, overlap-1))
        print(f"[INFO] Rechunk -> {len(chunks)} (size={size}, overlap={max(0, overlap-1)})")

    # If we got too few chunks (e.g., 1) and text is long, shrink a bit
    if len(chunks) <= 2 and len(text) > 12_000 and size > 2000:
        size = int(size * 0.85)
        size = max(1800, size)
        chunks = join_chunks(size, text, overlap_sents=min(2, overlap+1))
        print(f"[INFO] Rechunk small -> {len(chunks)} (size={size}, overlap={min(2, overlap+1)})")

    # Hard cap
    MAX_CHUNKS = 60
    if len(chunks) > MAX_CHUNKS:
        print(f"[WARN] Limiting chunks from {len(chunks)} to {MAX_CHUNKS}.")
        chunks = chunks[:MAX_CHUNKS]


    chunk_counter = 0
    ans_bank = []
    emit_idx = 0

    # === outputs ===
    out_dir = "runs"
    os.makedirs(out_dir, exist_ok=True)
    stamp = time.strftime("%Y%m%d-%H%M%S")
    jsonl_path = os.path.join(out_dir, f"qae-{stamp}.jsonl")
    csv_path   = os.path.join(out_dir, f"qae-{stamp}.csv")

    # create files on first write
    if not os.path.exists(jsonl_path):
        open(jsonl_path, "w").close()
    if not os.path.exists(csv_path):
        with open(csv_path, "w", newline="") as f:
            csv.writer(f).writerow(["q_idx", "chunk_idx", "anchor", "Q", "A", "E"])

    # de-dupe state
    _seen_norm_qs: set[str] = set()
    _seen_token_sets: list[set[str]] = []

    # coverage
    coverage = {
        "total_chars": len(text),
        "chunks_count": len(chunks),
        "anchors_attempted": 0,
        "qa_ok": 0,
        "qa_fallback": 0,
        "unique_evidence": set(),
    }

    for chunk_idx, chunk_raw in enumerate(chunks, start=1):
        chunk_counter = chunk_idx
        kw_raw = extract_keywords(chunk_raw)
        kw_8 = clean_anchors(kw_raw, chunk_raw, want_n=12)
        kw_8 = dedupe_anchors(kw_8)[:8]
        print(f"[DEBUG] anchors for chunk {chunk_idx}: {kw_8[:6]}")

        alpha_ratio = sum(ch.isalpha() for ch in chunk_raw) / max(1, len(chunk_raw))
        if alpha_ratio < 0.35 or len(kw_8) == 0:
            print(f"[SKIP] chunk {chunk_idx}: low alpha ({alpha_ratio:.2f}) or no anchors")
            continue

        has_ok_anchors = len([kw for kw in kw_8 if len(kw) >= 4 and re.search(r'[a-z]', kw, re.I)]) >= 3
        top_k = 1
        if size >= 2400 and has_ok_anchors and alpha_ratio >= 0.45:
            top_k = 2

        anchors_tried = 0

        for anchor_raw in kw_8[:top_k]:
            coverage["anchors_attempted"] += 1

            response = llm_generate_questions(chunk_raw, anchor_raw, q=None, temperature=0.2)
            item = parse_qae(response)

            if not item['ok']:
                response = llm_generate_questions(chunk_raw, anchor_raw, q=item.get('q') or "", temperature=0.2)
                item = parse_qae(response)
                if not item['ok']:
                    # <<< FIX: natural How/Why fallback here too (no “Explain … what it is …”)
                    sents = re.split(r'(?<=[.!?])\s+', chunk_raw.strip())
                    def _norm(s): return normalize_for_match(s)
                    e_fallback = next(
                        (s for s in sents if _norm(anchor_raw) in _norm(s)),
                        sents[0] if sents else chunk_raw.strip()
                    )
                    if not re.search(r'[.!?]"?$', e_fallback):
                        e_fallback += '.'
                    q_fallback = f"How does {anchor_raw} affect the topic discussed here, and why is it important?"
                    a_fallback = ("It shapes the ideas presented by influencing the structure/behavior being described. "
                                  "Its role leads to specific outcomes in the scenario, which explains why it matters in this part of the text.")
                    item = {'ok': True, 'q': q_fallback, 'a': a_fallback, 'e': e_fallback, 'raw': '', 'missing': set()}
                    coverage["qa_fallback"] += 1

            q = item['q']; a = item['a']; evidence = item['e']

            # sanitize anchor and skip junky ones
            anchor_raw = clean_text(anchor_raw).strip().strip('",.- ')
            if not re.match(r"[A-Za-z]", anchor_raw):
                words = re.findall(r"[A-Za-z]{3,}", chunk_raw)
                anchor_raw = " ".join(words[:2]) if words else "key idea"
            if (not re.search(r'[A-Za-z]', anchor_raw)) or len(anchor_raw) <= 2:
                m = re.search(r'(?i)\b([A-Za-z][A-Za-z\- ]{3,40})\b', q or '')
                anchor_raw = (m.group(1).strip() if m else 'key idea')

            # ensure evidence is a clean, single sentence that ends properly
            if evidence:
                evidence = evidence.strip()
                if not re.search(r'[.!?]"?$', evidence):
                    evidence += '.'
                if evidence not in chunk_raw:
                    sents = re.split(r'(?<=[.!?])\s+', chunk_raw.strip())
                    norm_anchor = normalize_for_match(anchor_raw)
                    cand = next((s for s in sents if norm_anchor in normalize_for_match(s)), None)
                    if cand:
                        evidence = cand.strip()
                        if not re.search(r'[.!?]"?$', evidence):
                            evidence += '.'

            # verify evidence + anchor; may regenerate once
            chunk_norm, _, _ = validate_qs(chunk_raw, kw_8, evidence)
            anchor_norm = normalize_for_match(anchor_raw)
            q, a, evidence = check_evidence_chunk(
                chunk_norm=chunk_norm, chunk_raw=chunk_raw,
                anchor_norm=anchor_norm, anchor_raw=anchor_raw,
                e_raw=evidence, q_raw=q, a_raw=a, kw_8=kw_8, regen_left=1
            )

            # coverage
            coverage["qa_ok"] += 1
            if evidence:
                coverage["unique_evidence"].add(evidence)

            # outputs
            ans_bank.append(a)
            anchors_tried += 1
            emit_idx += 1

            # print to console
            print(f"{chunk_counter}.{anchors_tried}")
            print(q)

            # append to files
            with open(jsonl_path, "a", encoding="utf-8") as jf:
                jf.write(json.dumps({
                    "q_idx": emit_idx,
                    "chunk_idx": chunk_counter,
                    "anchor": anchor_raw,
                    "Q": q, "A": a, "E": evidence
                }, ensure_ascii=False) + "\n")
            with open(csv_path, "a", newline="", encoding="utf-8") as cf:
                csv.writer(cf).writerow([emit_idx, chunk_counter, anchor_raw, q, a, evidence])

    # --- build Quizlet-ready exports ---
    quizlet = input("Would you like to export Q/A for quizlet?: ").lower()
    if quizlet in ["y", "yes"]:
        quizlet_dir = out_dir
        qa_path = os.path.join(quizlet_dir, f"quizlet-qa-{stamp}.csv")   # Q -> A
        qe_path = os.path.join(quizlet_dir, f"quizlet-qe-{stamp}.csv")   # Q -> E

        rows = []
        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            r = csv.DictReader(f)
            for row in r:
                Q = (row.get("Q") or "").replace("\n", " ").strip()
                A = (row.get("A") or "").replace("\n", " ").strip()
                E = (row.get("E") or "").replace("\n", " ").strip()
                rows.append((Q, A, E))

        # write Q->A
        with open(qa_path, "w", encoding="utf-8", newline="") as f:
            w = csv.writer(f)
            w.writerow(["Question", "Answer"])
            for Q, A, _E in rows:
                if Q and A:
                    w.writerow([Q, A])

        # write Q->E
        with open(qe_path, "w", encoding="utf-8", newline="") as f:
            w = csv.writer(f)
            w.writerow(["Question", "Answer"])
            for Q, _A, E in rows:
                if Q and E:
                    w.writerow([Q, E])

        print(f"[OK] Quizlet files written:\n - {qa_path}\n - {qe_path}")

    print("\n--- run summary ---")
    print(f"chars: {coverage['total_chars']}")
    print(f"chunks: {coverage['chunks_count']}")
    print(f"anchors attempted: {coverage['anchors_attempted']}")
    print(f"Q/A ok: {coverage['qa_ok']}  |  fallbacks: {coverage['qa_fallback']}")
    print(f"unique evidence sentences: {len(coverage['unique_evidence'])}")

    answer = input("Would you like to see answers? :").lower()
    if answer in ["y", "yes"]:
        for idx, ans in enumerate(ans_bank, start=1):
            print(f"{idx}) {ans}\n")


if __name__ == "__main__":
    # IMPORTANT for multiprocessing on macOS/Windows and for safety on Linux
    try:
        mp.set_start_method("spawn", force=True)
    except RuntimeError:
        pass

    main()

